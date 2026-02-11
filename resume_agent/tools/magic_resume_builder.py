#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Magic Resume JSON Builder - 将简历内容转换为 Magic Resume JSON 格式
并支持导出为 PDF 和 Word 文档
"""

from __future__ import annotations

import json
import re
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict, Any

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


@dataclass
class StyleSettings:
    """样式配置类，从 globalSettings 读取"""
    base_font_size: int = 16
    page_padding: int = 32
    paragraph_spacing: int = 12
    line_height: float = 1.3
    section_spacing: int = 10
    header_size: int = 18
    subheader_size: int = 16
    theme_color: str = "#000000"
    center_subtitle: bool = True
    
    @classmethod
    def from_dict(cls, d: dict) -> "StyleSettings":
        return cls(
            base_font_size=d.get("baseFontSize", 16),
            page_padding=d.get("pagePadding", 32),
            paragraph_spacing=d.get("paragraphSpacing", 12),
            line_height=d.get("lineHeight", 1.3),
            section_spacing=d.get("sectionSpacing", 10),
            header_size=d.get("headerSize", 18),
            subheader_size=d.get("subheaderSize", 16),
            theme_color=d.get("themeColor", "#000000"),
            center_subtitle=d.get("centerSubtitle", True),
        )
    
    def px_to_pt(self, px: int) -> float:
        """px 转 pt"""
        return px * 0.58
    
    def get_color(self) -> RGBColor:
        """解析主题色"""
        try:
            r = int(self.theme_color[1:3], 16)
            g = int(self.theme_color[3:5], 16)
            b = int(self.theme_color[5:7], 16)
            return RGBColor(r, g, b)
        except (ValueError, IndexError):
            return RGBColor(0, 0, 0)


@dataclass
class ResumeData:
    """简历数据结构"""
    # 基本信息
    name: str = ""
    title: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    
    # 教育经历
    education: List[Dict[str, Any]] = field(default_factory=list)
    
    # 工作经历
    experience: List[Dict[str, Any]] = field(default_factory=list)
    
    # 项目经历
    projects: List[Dict[str, Any]] = field(default_factory=list)
    
    # 技能
    skill_content: str = ""
    
    # 个人简介
    summary: str = ""


def strip_html(html_content: str) -> str:
    """去除 HTML 标签，保留纯文本"""
    if not html_content:
        return ""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")
        return soup.get_text(separator="\n").strip()
    except ImportError:
        # 简单的 HTML 标签移除
        return re.sub(r'<[^>]+>', '', html_content).strip()


def html_to_lines(html_content: str) -> list:
    """将 HTML 内容转为行列表，保留列表项结构"""
    if not html_content:
        return []
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")
        lines = []
        for li in soup.find_all("li"):
            text = li.get_text(strip=True)
            if text:
                lines.append(text)
        if not lines:
            text = soup.get_text(separator="\n").strip()
            lines = [line.strip() for line in text.split("\n") if line.strip()]
        return lines
    except ImportError:
        # 简单处理
        text = re.sub(r'<[^>]+>', '\n', html_content)
        return [line.strip() for line in text.split("\n") if line.strip()]


def markdown_to_html(markdown_text: str) -> str:
    """将 Markdown 转换为 HTML"""
    import re
    
    def convert_inline_markdown(text: str) -> str:
        """转换行内 Markdown 语法"""
        # 粗体: **text** 或 __text__
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'__(.+?)__', r'<strong>\1</strong>', text)
        # 斜体: *text* 或 _text_
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        text = re.sub(r'_(.+?)_', r'<em>\1</em>', text)
        # 行内代码: `code`
        text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
        return text
    
    lines = markdown_text.split("\n")
    html_parts = []
    in_list = False
    
    for line in lines:
        line = line.strip()
        if not line:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            continue
        
        # 转换行内 Markdown
        line = convert_inline_markdown(line)
        
        # 处理列表项
        if line.startswith("- ") or line.startswith("* "):
            if not in_list:
                html_parts.append('<ul class="custom-list">')
                in_list = True
            content = line[2:].strip()
            html_parts.append(f"<li><p>{content}</p></li>")
        else:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<p>{line}</p>")
    
    if in_list:
        html_parts.append("</ul>")
    
    return "\n".join(html_parts)


class MagicResumeBuilder:
    """
    Magic Resume JSON 构建器
    将 TextModifier 生成的各模块内容填充到 Magic Resume JSON 格式
    
    支持的模板:
    - classic: 经典模板（默认）- 传统简约的简历布局
    - modern: 两栏布局 - 突出个人特色
    - left-right: 模块标题背景色 - 美观突出
    - timeline: 时间线风格 - 突出时间顺序
    """
    
    # 可用模板配置
    TEMPLATES = {
        "classic": {
            "name": "经典模板",
            "description": "传统简约的简历布局，适合大多数求职场景",
            "layout": "classic",
            "colorScheme": {"primary": "#000000", "secondary": "#4b5563", "background": "#ffffff", "text": "#212529"},
            "spacing": {"sectionGap": 24, "itemGap": 16, "contentPadding": 32},
            "basic": {"layout": "center"}
        },
        "modern": {
            "name": "两栏布局",
            "description": "经典两栏，突出个人特色",
            "layout": "modern",
            "colorScheme": {"primary": "#000000", "secondary": "#6b7280", "background": "#ffffff", "text": "#212529"},
            "spacing": {"sectionGap": 20, "itemGap": 20, "contentPadding": 1},
            "basic": {"layout": "center"}
        },
        "left-right": {
            "name": "模块标题背景色",
            "description": "模块标题背景鲜明，突出美观特色",
            "layout": "left-right",
            "colorScheme": {"primary": "#000000", "secondary": "#9ca3af", "background": "#ffffff", "text": "#212529"},
            "spacing": {"sectionGap": 24, "itemGap": 16, "contentPadding": 32},
            "basic": {"layout": "left"}
        },
        "timeline": {
            "name": "时间线风格",
            "description": "时间线布局，突出经历的时间顺序",
            "layout": "timeline",
            "colorScheme": {"primary": "#18181b", "secondary": "#64748b", "background": "#ffffff", "text": "#212529"},
            "spacing": {"sectionGap": 1, "itemGap": 12, "contentPadding": 24},
            "basic": {"layout": "right"}
        }
    }
    
    DEFAULT_GLOBAL_SETTINGS = {
        "baseFontSize": 16,
        "pagePadding": 32,
        "paragraphSpacing": 12,
        "lineHeight": 1.3,
        "sectionSpacing": 10,
        "headerSize": 18,
        "subheaderSize": 16,
        "useIconMode": True,
        "themeColor": "#000000",
        "centerSubtitle": True
    }
    
    DEFAULT_MENU_SECTIONS = [
        {"id": "basic", "title": "基本信息", "icon": "👤", "enabled": True, "order": 0},
        {"id": "education", "title": "教育经历", "icon": "🎓", "enabled": True, "order": 1},
        {"id": "experience", "title": "工作经验", "icon": "💼", "enabled": True, "order": 2},
        {"id": "projects", "title": "项目经历", "icon": "🚀", "enabled": True, "order": 3},
        {"id": "skills", "title": "专业技能", "icon": "⚡", "enabled": True, "order": 4},
    ]

    def __init__(self, template_id: str = "classic"):
        """
        初始化构建器
        
        Args:
            template_id: 模板ID，可选值: classic, modern, left-right, timeline
        """
        if template_id not in self.TEMPLATES:
            print(f"警告: 未知模板 '{template_id}'，使用默认模板 'classic'")
            template_id = "classic"
        
        self.template_id = template_id
        self.template_config = self.TEMPLATES[template_id]
        self.resume_id = str(uuid.uuid4())
        
        # 合并模板配置到全局设置
        global_settings = self.DEFAULT_GLOBAL_SETTINGS.copy()
        global_settings["themeColor"] = self.template_config["colorScheme"]["primary"]
        
        self.data: Dict[str, Any] = {
            "title": f"简历_{self.resume_id[:8]}",
            "id": self.resume_id,
            "createdAt": datetime.now().isoformat(),
            "updatedAt": datetime.now().isoformat(),
            "templateId": template_id,
            "basic": {},
            "education": [],
            "experience": [],
            "projects": [],
            "skillContent": "",
            "menuSections": self.DEFAULT_MENU_SECTIONS.copy(),
            "globalSettings": global_settings,
            "customData": {},
        }
    
    @classmethod
    def list_templates(cls) -> Dict[str, str]:
        """列出所有可用模板"""
        return {tid: t["name"] for tid, t in cls.TEMPLATES.items()}
    
    def set_basic_info(
        self, 
        name: str = "",
        title: str = "",
        email: str = "",
        phone: str = "",
        location: str = "",
        **kwargs
    ) -> "MagicResumeBuilder":
        """设置基本信息"""
        self.data["basic"] = {
            "name": name,
            "title": title,
            "email": email,
            "phone": phone,
            "location": location,
            "fieldOrder": [
                {"id": "1", "key": "name", "label": "姓名", "type": "text", "visible": True},
                {"id": "2", "key": "title", "label": "职位", "type": "text", "visible": bool(title)},
                {"id": "5", "key": "email", "label": "邮箱", "type": "text", "visible": bool(email)},
                {"id": "6", "key": "phone", "label": "电话", "type": "text", "visible": bool(phone)},
                {"id": "7", "key": "location", "label": "所在地", "type": "text", "visible": bool(location)},
            ],
            "icons": {
                "email": "Mail",
                "phone": "Phone",
                "location": "MapPin"
            },
            "photoConfig": {
                "width": 90,
                "height": 120,
                "aspectRatio": "1:1",
                "borderRadius": "none",
                "visible": False
            },
            "customFields": [],
            **kwargs
        }
        return self
    
    def add_education(
        self,
        school: str,
        major: str = "",
        degree: str = "",
        start_date: str = "",
        end_date: str = "",
        gpa: str = "",
        description: str = ""
    ) -> "MagicResumeBuilder":
        """添加教育经历"""
        edu_id = str(uuid.uuid4())
        self.data["education"].append({
            "id": edu_id,
            "school": school,
            "major": major,
            "degree": degree,
            "startDate": start_date,
            "endDate": end_date,
            "gpa": gpa,
            "description": markdown_to_html(description) if description else "",
            "visible": True
        })
        return self
    
    def add_experience(
        self,
        company: str,
        position: str = "",
        date: str = "",
        details: str = ""
    ) -> "MagicResumeBuilder":
        """添加工作经历"""
        exp_id = str(uuid.uuid4())
        self.data["experience"].append({
            "id": exp_id,
            "company": company,
            "position": position,
            "date": date,
            "details": markdown_to_html(details) if details else "",
            "visible": True
        })
        return self
    
    def add_project(
        self,
        name: str,
        role: str = "",
        date: str = "",
        description: str = "",
        link: str = ""
    ) -> "MagicResumeBuilder":
        """添加项目经历"""
        proj_id = str(uuid.uuid4())
        self.data["projects"].append({
            "id": proj_id,
            "name": name,
            "role": role,
            "date": date,
            "description": markdown_to_html(description) if description else "",
            "link": link,
            "visible": True
        })
        return self
    
    def set_skills(self, skill_content: str) -> "MagicResumeBuilder":
        """设置技能内容"""
        self.data["skillContent"] = markdown_to_html(skill_content)
        return self
    
    def set_global_settings(self, **kwargs) -> "MagicResumeBuilder":
        """设置全局样式"""
        self.data["globalSettings"].update(kwargs)
        return self
    
    def build(self) -> Dict[str, Any]:
        """构建并返回 JSON 数据"""
        self.data["updatedAt"] = datetime.now().isoformat()
        return self.data
    
    def to_json(self, output_path: Path) -> Path:
        """导出为 JSON 文件"""
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(self.build(), f, ensure_ascii=False, indent=2)
        return output_path


class MagicResumeDocxBuilder:
    """将 Magic Resume JSON 转换为 Word 文档"""
    
    def __init__(self, data: dict):
        self.data = data
        self.settings = StyleSettings.from_dict(data.get("globalSettings", {}))
        self.doc = Document()
        self.color = self.settings.get_color()
        self._setup_document()
    
    def _setup_document(self):
        """设置文档基本属性"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
        
        style = self.doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(self.settings.px_to_pt(self.settings.base_font_size))
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def _add_paragraph(self, text: str, font_size: float, bold: bool = False, 
                       color: RGBColor = None, alignment: int = WD_ALIGN_PARAGRAPH.LEFT,
                       space_before: float = 0, space_after: float = 0,
                       line_spacing: float = None, indent: float = 0):
        """添加段落"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if color:
            run.font.color.rgb = color
        
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        
        if line_spacing:
            p.paragraph_format.line_spacing = line_spacing
        
        return p
    
    def _add_section_title(self, title: str):
        """添加章节标题"""
        font_size = self.settings.px_to_pt(self.settings.header_size)
        p = self._add_paragraph(
            title,
            font_size=font_size,
            bold=True,
            color=self.color,
            space_before=4,
            space_after=1,
        )
        # 添加底部边框
        self._add_bottom_border(p)
        return p
    
    def _add_bottom_border(self, paragraph):
        """为段落添加底部边框线"""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), self.settings.theme_color.replace('#', ''))
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    def _add_item_header(self, left_text: str, right_text: str = ""):
        """添加条目标题行"""
        font_size = self.settings.px_to_pt(self.settings.subheader_size)
        
        if right_text:
            table = self.doc.add_table(rows=1, cols=2)
            table.autofit = True
            
            left_cell = table.cell(0, 0)
            left_p = left_cell.paragraphs[0]
            left_run = left_p.add_run(left_text)
            left_run.font.size = Pt(font_size)
            left_run.font.bold = True
            left_run.font.name = '微软雅黑'
            left_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            right_cell = table.cell(0, 1)
            right_p = right_cell.paragraphs[0]
            right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_run = right_p.add_run(right_text)
            right_run.font.size = Pt(font_size * 0.85)
            right_run.font.name = '微软雅黑'
            right_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            right_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # 移除表格边框
            for row in table.rows:
                for cell in row.cells:
                    self._remove_cell_border(cell)
        else:
            self._add_paragraph(
                left_text,
                font_size=font_size,
                bold=True,
                space_before=3,
                space_after=0,
            )
    
    def _remove_cell_border(self, cell):
        """移除单元格边框"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'nil')
            tcBorders.append(element)
        tcPr.append(tcBorders)
    
    def _add_list_item(self, text: str):
        """添加列表项"""
        font_size = self.settings.px_to_pt(self.settings.base_font_size)
        return self._add_paragraph(
            f"•  {text}",
            font_size=font_size,
            space_after=0,
            line_spacing=1.0,
        )
    
    def add_basic_info(self, basic: dict):
        """添加基本信息"""
        name_size = self.settings.px_to_pt(self.settings.header_size) * 1.3
        
        name = basic.get("name", "")
        if name:
            self._add_paragraph(
                name,
                font_size=name_size,
                bold=True,
                color=self.color,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=1,
            )
        
        title = basic.get("title", "")
        if title:
            self._add_paragraph(
                title,
                font_size=self.settings.px_to_pt(self.settings.base_font_size),
                color=RGBColor(100, 100, 100),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=2,
            )
        
        # 联系方式
        contact_parts = []
        for key in ["email", "phone", "location"]:
            value = basic.get(key, "")
            if value:
                contact_parts.append(value)
        
        if contact_parts:
            self._add_paragraph(
                "  |  ".join(contact_parts),
                font_size=self.settings.px_to_pt(self.settings.base_font_size) * 0.85,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=4,
            )
    
    def add_education(self, education_list: list):
        """添加教育经历"""
        if not education_list:
            return
        
        self._add_section_title("教育经历")
        
        for edu in education_list:
            if not edu.get("visible", True):
                continue
            
            school = edu.get("school", "")
            degree = edu.get("degree", "")
            major = edu.get("major", "")
            
            header_parts = [p for p in [school, degree, major] if p]
            header = "  |  ".join(header_parts)
            
            date_str = edu.get("date", "") or ""
            if not date_str:
                start = edu.get("startDate", "")
                end = edu.get("endDate", "")
                if start or end:
                    date_str = f"{start[:7] if start else ''} - {end[:7] if end else ''}"
            
            self._add_item_header(header, date_str)
            
            gpa = edu.get("gpa", "")
            if gpa:
                self._add_paragraph(
                    f"GPA: {gpa}",
                    font_size=self.settings.px_to_pt(self.settings.base_font_size) * 0.85,
                    color=RGBColor(80, 80, 80),
                    space_after=0,
                )
            
            desc = edu.get("description", "")
            if desc:
                for line in html_to_lines(desc):
                    self._add_list_item(line)
    
    def add_experience(self, experience_list: list):
        """添加工作经验"""
        if not experience_list:
            return
        
        self._add_section_title("工作经验")
        
        for exp in experience_list:
            if not exp.get("visible", True):
                continue
            
            company = exp.get("company", "")
            position = exp.get("position", "")
            date_str = exp.get("date", "")
            
            header = f"{company}  |  {position}" if company and position else company or position
            self._add_item_header(header, date_str)
            
            details = exp.get("details", "")
            if details:
                for line in html_to_lines(details):
                    self._add_list_item(line)
    
    def add_projects(self, projects_list: list):
        """添加项目经历"""
        if not projects_list:
            return
        
        self._add_section_title("项目经历")
        
        for proj in projects_list:
            if not proj.get("visible", True):
                continue
            
            name = proj.get("name", "")
            role = proj.get("role", "")
            date_str = proj.get("date", "")
            
            header = f"{name}  |  {role}" if name and role else name or role
            self._add_item_header(header, date_str)
            
            desc = proj.get("description", "")
            if desc:
                for line in html_to_lines(desc):
                    self._add_list_item(line)
    
    def add_skills(self, skill_content: str):
        """添加专业技能"""
        if not skill_content:
            return
        
        self._add_section_title("专业技能")
        
        lines = html_to_lines(skill_content)
        if lines:
            for line in lines:
                self._add_list_item(line)
        else:
            text = strip_html(skill_content)
            if text:
                self._add_paragraph(
                    text,
                    font_size=self.settings.px_to_pt(self.settings.base_font_size),
                    space_after=4,
                )
    
    def build(self) -> Document:
        """构建文档"""
        menu_sections = self.data.get("menuSections", [])
        menu_sections = sorted(menu_sections, key=lambda x: x.get("order", 0))
        
        basic = self.data.get("basic", {})
        education = self.data.get("education", [])
        experience = self.data.get("experience", [])
        projects = self.data.get("projects", [])
        skill_content = self.data.get("skillContent", "")
        
        for section in menu_sections:
            if not section.get("enabled", True):
                continue
            
            section_id = section.get("id", "")
            
            if section_id == "basic":
                self.add_basic_info(basic)
            elif section_id == "education":
                self.add_education(education)
            elif section_id == "experience":
                self.add_experience(experience)
            elif section_id == "projects":
                self.add_projects(projects)
            elif section_id == "skills":
                self.add_skills(skill_content)
        
        return self.doc
    
    def save(self, output_path: str):
        """保存文档"""
        self.build()
        self.doc.save(output_path)
