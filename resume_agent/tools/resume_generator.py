from __future__ import annotations

import os
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Optional


@dataclass
class GenerateRequest:
    markdown_text: str
    output_dir: Path
    output_format: str
    candidate_name: str = "候选人"


class ResumeGenerator:
    """Generate resume PDF via LaTeX and optionally a Word doc."""

    def run(self, request: GenerateRequest) -> Path:
        request.output_dir.mkdir(parents=True, exist_ok=True)

        md_path = request.output_dir / "resume.md"
        md_path.write_text(request.markdown_text, encoding="utf-8")

        tex_path = request.output_dir / "resume.tex"
        tex_path.write_text(self._render_latex(request), encoding="utf-8")

        if request.output_format.lower() == "pdf":
            pdf_path = request.output_dir / "resume.pdf"
            self._compile_pdf(tex_path, request.output_dir)
            if not pdf_path.exists():
                raise RuntimeError("LaTeX compilation failed: resume.pdf not found.")
            return pdf_path

        if request.output_format.lower() in ("docx", "word"):
            docx_path = request.output_dir / "resume.docx"
            self._generate_docx(request.markdown_text, docx_path)
            return docx_path

        raise ValueError(f"Unsupported output format: {request.output_format}")

    def _render_latex(self, request: GenerateRequest) -> str:
        escaped = self._escape_latex(request.markdown_text)
        return f"""\\documentclass[11pt]{{article}}
\\usepackage[margin=1in]{{geometry}}
\\usepackage{{hyperref}}
\\usepackage{{titlesec}}
\\usepackage{{enumitem}}
\\titleformat{{\\section}}{{\\large\\bfseries}}{{}}{{0em}}{{}}
\\setlist[itemize]{{noitemsep, topsep=0pt}}
\\begin{{document}}
\\begin{{center}}
{{\\LARGE {request.candidate_name}}}\\\\
\\end{{center}}
\\vspace{{0.5em}}
{escaped}
\\end{{document}}
"""

    def _escape_latex(self, text: str) -> str:
        replacements = {
            "&": "\\&",
            "%": "\\%",
            "$": "\\$",
            "#": "\\#",
            "_": "\\_",
            "{": "\\{",
            "}": "\\}",
            "~": "\\textasciitilde{}",
            "^": "\\textasciicircum{}",
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        lines = []
        for line in text.splitlines():
            if line.startswith("#"):
                title = line.lstrip("#").strip()
                lines.append(f"\\section*{{{title}}}")
            elif line.startswith("-"):
                if not lines or not lines[-1].startswith("\\begin{itemize}"):
                    lines.append("\\begin{itemize}")
                lines.append(f"\\item {line[1:].strip()}")
            else:
                if lines and lines[-1].startswith("\\item"):
                    lines.append("\\end{itemize}")
                if line.strip():
                    lines.append(line + "\\\\")
        if lines and lines[-1].startswith("\\item"):
            lines.append("\\end{itemize}")
        return "\n".join(lines)

    def _compile_pdf(self, tex_path: Path, output_dir: Path) -> None:
        try:
            subprocess.run(
                ["pdflatex", "-interaction=nonstopmode", tex_path.name],
                cwd=str(output_dir),
                check=False,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
        except FileNotFoundError as exc:
            raise RuntimeError("pdflatex is required to generate PDF.") from exc

    def _generate_docx(self, markdown_text: str, output_path: Path) -> None:
        try:
            import docx  # type: ignore
        except Exception as exc:
            raise RuntimeError("python-docx is required to generate Word files.") from exc

        doc = docx.Document()
        for line in markdown_text.splitlines():
            if line.startswith("#"):
                doc.add_heading(line.lstrip("#").strip(), level=1)
            elif line.startswith("-"):
                doc.add_paragraph(line[1:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(line)
        doc.save(str(output_path))
